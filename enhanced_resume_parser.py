#!/usr/bin/env python3
"""
Enhanced Resume Parser - 增强版简历解析器
配置路径: enhanced_resume_parser.py (项目根目录)
"""
import os
import re
import json
import io
import fitz  # PyMuPDF
import pdfplumber
import pytesseract
from PIL import Image
import cv2
import numpy as np
from typing import Dict, Any, List, Optional
from dataclasses import dataclass, asdict
from loguru import logger
import docx
from datetime import datetime
import openai
import asyncio

# Configure logger
logger.remove()  # Remove default handler
logger.add("resume_parser.log", level="INFO", rotation="10 MB")
logger.add(lambda msg: print(msg, end=""), level="INFO")

@dataclass
class EnhancedPersonalInfo:
    """增强版个人信息"""
    full_name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    github: str = ""
    portfolio: str = ""
    years_of_experience: str = ""

@dataclass
class EnhancedJob:
    """增强版工作经历"""
    job_title: str = ""
    company: str = ""
    location: str = ""
    start_date: str = ""
    end_date: str = ""
    duration: str = ""
    responsibilities: List[str] = None
    achievements: List[str] = None
    technologies_used: List[str] = None
    
    def __post_init__(self):
        if self.responsibilities is None:
            self.responsibilities = []
        if self.achievements is None:
            self.achievements = []
        if self.technologies_used is None:
            self.technologies_used = []

@dataclass
class EnhancedEducation:
    """增强版教育背景"""
    school: str = ""
    degree: str = ""
    major: str = ""
    start_date: str = ""
    end_date: str = ""
    gpa: str = ""
    honors: str = ""
    relevant_courses: List[str] = None
    
    def __post_init__(self):
        if self.relevant_courses is None:
            self.relevant_courses = []

@dataclass
class TechnicalSkills:
    """技术技能分类"""
    programming_languages: List[str] = None
    frameworks_libraries: List[str] = None
    databases: List[str] = None
    cloud_platforms: List[str] = None
    tools_software: List[str] = None
    methodologies: List[str] = None
    soft_skills: List[str] = None
    
    def __post_init__(self):
        for field in ['programming_languages', 'frameworks_libraries', 'databases',
                     'cloud_platforms', 'tools_software', 'methodologies', 'soft_skills']:
            if getattr(self, field) is None:
                setattr(self, field, [])

@dataclass
class ProjectExperience:
    """项目经验"""
    name: str = ""
    description: str = ""
    role: str = ""
    duration: str = ""
    team_size: str = ""
    technologies: List[str] = None
    achievements: List[str] = None
    github_url: str = ""
    demo_url: str = ""
    
    def __post_init__(self):
        if self.technologies is None:
            self.technologies = []
        if self.achievements is None:
            self.achievements = []

@dataclass
class EnhancedResume:
    """增强版简历数据结构"""
    personal_info: EnhancedPersonalInfo
    technical_skills: TechnicalSkills
    work_experience: List[EnhancedJob]
    education: List[EnhancedEducation]
    projects: List[ProjectExperience]
    certifications: List[str]
    notable_achievements: List[str]
    languages: List[str]
    raw_text: str = ""
    parsing_confidence: float = 0.0
    
    def __post_init__(self):
        if not self.work_experience:
            self.work_experience = []
        if not self.education:
            self.education = []
        if not self.projects:
            self.projects = []
        if not self.certifications:
            self.certifications = []
        if not self.notable_achievements:
            self.notable_achievements = []
        if not self.languages:
            self.languages = []

class EnhancedResumeParser:
    """增强版简历解析器"""
    
    def __init__(self, openai_api_key: str):
        self.openai_client = openai.AsyncOpenAI(
            api_key=openai_api_key,
            timeout=30.0,
            max_retries=3
        )
        
        # 技术关键词词典 - 扩展版
        self.tech_keywords = {
            'programming_languages': [
                'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'c', 'go', 'rust',
                'php', 'ruby', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql', 'html',
                'css', 'sass', 'less', 'shell', 'bash', 'powershell', 'perl', 'dart', 'lua'
            ],
            'frameworks_libraries': [
                'react', 'vue', 'angular', 'django', 'flask', 'fastapi', 'spring', 'nodejs',
                'express', 'nextjs', 'nuxtjs', 'laravel', 'rails', 'asp.net', 'tensorflow',
                'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy', 'opencv', 'jquery',
                'bootstrap', 'tailwind', 'material-ui', 'ant-design', 'element-ui'
            ],
            'databases': [
                'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'oracle',
                'sqlite', 'cassandra', 'dynamodb', 'neo4j', 'influxdb', 'mariadb',
                'couchdb', 'firebase', 'supabase'
            ],
            'cloud_platforms': [
                'aws', 'azure', 'gcp', 'google cloud', 'alibaba cloud', 'tencent cloud',
                'heroku', 'vercel', 'netlify', 'digitalocean', 'linode', 'vultr'
            ],
            'tools_software': [
                'docker', 'kubernetes', 'git', 'github', 'gitlab', 'jenkins', 'gitlab-ci',
                'github-actions', 'ansible', 'terraform', 'vagrant', 'jira', 'confluence',
                'slack', 'notion', 'figma', 'sketch', 'photoshop', 'illustrator', 'xd',
                'postman', 'insomnia', 'swagger', 'linux', 'ubuntu', 'centos', 'windows'
            ],
            'methodologies': [
                'agile', 'scrum', 'kanban', 'devops', 'ci/cd', 'tdd', 'bdd', 'microservices',
                'restful', 'graphql', 'api', 'mvc', 'mvvm', 'solid', 'design patterns',
                'clean architecture', 'domain-driven design'
            ]
        }
        
        # 软技能关键词
        self.soft_skills_keywords = [
            'leadership', 'teamwork', 'communication', 'problem solving', 'critical thinking',
            'project management', 'time management', 'adaptability', 'creativity', 'innovation',
            'analytical thinking', 'decision making', 'negotiation', 'presentation', 'mentoring',
            'collaboration', 'cross-functional', 'stakeholder management', 'customer service'
        ]
    
    async def parse_resume(self, file_path: str = None, text: str = None) -> EnhancedResume:
        """解析简历 - 主入口"""
        try:
            if file_path:
                text = await self._extract_text_enhanced(file_path)
            elif not text:
                raise ValueError("必须提供文件路径或文本内容")
            
            logger.info(f"提取的文本长度: {len(text)} 字符")
            
            # 文本预处理
            cleaned_text = self._preprocess_text(text)
            
            # 使用增强版OpenAI提取
            structured_data = await self._extract_with_enhanced_openai(cleaned_text)
            
            # 正则表达式辅助提取技术技能
            regex_skills = self._extract_skills_with_regex(cleaned_text)
            
            # 合并和验证结果
            final_data = self._merge_and_validate_results(structured_data, regex_skills)
            
            # 计算解析置信度
            confidence = self._calculate_parsing_confidence(final_data, cleaned_text)
            
            # 构建增强版简历对象
            resume = self._build_enhanced_resume(final_data, cleaned_text, confidence)
            
            logger.info(f"简历解析完成，置信度: {confidence:.2f}")
            return resume
            
        except Exception as e:
            logger.error(f"简历解析失败: {e}")
            return self._create_fallback_resume(text or "")
    
    async def _extract_text_enhanced(self, file_path: str) -> str:
        """增强版文本提取，支持PDF、DOCX、TXT格式"""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            extension = file_path.lower().split('.')[-1]
            
            if extension == 'pdf':
                return await self._extract_from_pdf_enhanced(file_path)
            elif extension in ['docx', 'doc']:
                return self._extract_from_docx_enhanced(file_path)
            elif extension == 'txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            else:
                raise ValueError(f"不支持的文件格式: {extension}")
        
        except Exception as e:
            logger.error(f"文本提取失败: {e}")
            return ""
    
    async def _extract_from_pdf_enhanced(self, file_path: str) -> str:
        """增强版PDF文本提取，结合pdfplumber和OCR"""
        text = ""
        
        try:
            # 方法1: 使用pdfplumber
            logger.info("尝试使用pdfplumber提取PDF文本...")
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # 如果提取的文本太少，尝试OCR
            if len(text.strip()) < 100:
                logger.info("文本提取不足，尝试OCR识别...")
                text = await self._extract_with_ocr(file_path)
            
        except Exception as e:
            logger.warning(f"pdfplumber提取失败: {e}，尝试PyMuPDF...")
            try:
                # 方法2: 使用PyMuPDF
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
                
                if len(text.strip()) < 100:
                    text = await self._extract_with_ocr(file_path)
                    
            except Exception as e2:
                logger.error(f"PyMuPDF提取失败: {e2}，尝试OCR...")
                text = await self._extract_with_ocr(file_path)
        
        return text
    
    async def _extract_with_ocr(self, file_path: str) -> str:
        """OCR图片文字识别，使用pytesseract"""
        try:
            logger.info("开始OCR识别...")
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 提高分辨率
                img_data = pix.pil_tobytes(format="PNG")
                
                # 使用PIL打开图片
                image = Image.open(io.BytesIO(img_data))
                
                # 图片预处理
                image = self._preprocess_image_for_ocr(image)
                
                # OCR识别，支持英文和简体中文
                page_text = pytesseract.image_to_string(image, lang='eng+chi_sim')
                text += page_text + "\n"
            
            doc.close()
            logger.info(f"OCR识别完成，提取文本长度: {len(text)}")
            return text
            
        except pytesseract.TesseractNotFoundError:
            logger.error("Tesseract OCR未安装或未正确配置")
            return "OCR提取失败：请确保Tesseract OCR已安装"
        except Exception as e:
            logger.error(f"OCR识别失败: {e}")
            return "OCR提取失败：请检查文件格式或pytesseract配置"
    
    def _preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """图片预处理以提高OCR准确率"""
        try:
            img_array = np.array(image)
            
            # 转换为灰度图
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
            
            # 去噪
            denoised = cv2.medianBlur(gray, 3)
            
            # 二值化
            _, binary = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            return Image.fromarray(binary)
            
        except Exception as e:
            logger.error(f"图片预处理失败: {e}")
            return image
    
    def _extract_from_docx_enhanced(self, file_path: str) -> str:
        """增强版DOCX文本提取"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # 提取段落
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"DOCX提取失败: {e}")
            return ""
    
    def _preprocess_text(self, text: str) -> str:
        """文本预处理，规范化编码和格式"""
        try:
            # 确保UTF-8编码
            text = text.encode('utf-8', errors='ignore').decode('utf-8')
            
            # 移除多余空白
            text = re.sub(r'\s+', ' ', text)
            
            # 移除无用特殊字符，保留标点
            text = re.sub(r'[^\w\s\-@.(),/]', ' ', text)
            
            # 标准化邮箱和电话格式
            text = re.sub(r'\b([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b', r' \1 ', text)
            text = re.sub(r'\b(\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})\b', r' \1 ', text)
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"文本预处理失败: {e}")
            return text
    
    async def _extract_with_enhanced_openai(self, text: str) -> Dict[str, Any]:
        """使用业界最佳实践的简历提取prompt"""
        
        # 检测是否包含中文
        is_chinese = self._contains_chinese(text)
        
        if is_chinese:
            prompt = f"""
    请从以下简历文本中提取结构化信息。严格按照以下YAML格式输出，不要添加任何其他文字：

    ---
    personal_info:
    full_name: '姓名'
    email: '邮箱地址'
    phone: '电话号码'
    location: '工作地点'
    linkedin: 'LinkedIn链接'
    github: 'GitHub链接'
    years_of_experience: '工作年限'

    technical_skills:
    programming_languages: ['Python', 'Java', 'JavaScript']
    frameworks_libraries: ['React', 'Django', 'Spring']
    databases: ['MySQL', 'MongoDB', 'Redis']
    cloud_platforms: ['AWS', 'Azure', 'GCP']
    tools_software: ['Docker', 'Git', 'Jenkins']

    work_experience:
    - job_title: '职位名称'
        company: '公司名称'
        start_date: '开始时间'
        end_date: '结束时间'
        responsibilities: ['主要职责1', '主要职责2']
        achievements: ['主要成就1', '主要成就2']

    education:
    - school: '学校名称'
        degree: '学位'
        major: '专业'
        start_date: '入学时间'
        end_date: '毕业时间'

    projects:
    - name: '项目名称'
        description: '项目描述'
        technologies: ['使用技术1', '使用技术2']
        achievements: ['项目成果1', '项目成果2']

    certifications: ['证书1', '证书2']
    languages: ['语言能力1', '语言能力2']

    简历文本：
    {text}

    重要：只输出YAML格式，从---开始，仔细提取所有个人信息、工作经历和教育背景。
    """
        else:
            prompt = f"""
    Extract structured information from this resume text. Output ONLY in YAML format, no other text:

    ---
    personal_info:
    full_name: 'Full Name'
    email: 'email@example.com'
    phone: 'Phone Number'
    location: 'Location'
    linkedin: 'LinkedIn URL'
    github: 'GitHub URL'
    years_of_experience: 'Years'

    technical_skills:
    programming_languages: ['Python', 'Java', 'JavaScript']
    frameworks_libraries: ['React', 'Django', 'Spring']
    databases: ['MySQL', 'MongoDB', 'Redis']
    cloud_platforms: ['AWS', 'Azure', 'GCP']
    tools_software: ['Docker', 'Git', 'Jenkins']

    work_experience:
    - job_title: 'Job Title'
        company: 'Company Name'
        start_date: 'Start Date'
        end_date: 'End Date'
        responsibilities: ['Responsibility 1', 'Responsibility 2']
        achievements: ['Achievement 1', 'Achievement 2']

    education:
    - school: 'School Name'
        degree: 'Degree'
        major: 'Major'
        start_date: 'Start Date'
        end_date: 'End Date'

    projects:
    - name: 'Project Name'
        description: 'Project Description'
        technologies: ['Tech 1', 'Tech 2']
        achievements: ['Achievement 1', 'Achievement 2']

    certifications: ['Certification 1', 'Certification 2']
    languages: ['Language 1', 'Language 2']

    Resume text:
    {text}

    Important: Output ONLY YAML format starting with ---, extract ALL personal info, work experience, and education carefully.
    """

        try:
            # 使用gpt-3.5-turbo，温度设为0
            response = await self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",  # 改用3.5，更适合结构化任务
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0  # 设为0确保一致性
            )
            
            result = response.choices[0].message.content.strip()
            
            # 解析YAML
            import yaml
            parsed_data = yaml.safe_load(result)
            
            logger.info("YAML解析成功")
            logger.info(f"提取的个人信息: {parsed_data.get('personal_info', {})}")
            
            return parsed_data
            
        except yaml.YAMLError as e:
            logger.error(f"YAML解析失败: {e}")
            logger.error(f"原始响应: {result[:300]}...")
            return self._create_fallback_structure()
            
        except Exception as e:
            logger.error(f"OpenAI提取失败: {e}")
            return self._create_fallback_structure()

    def _contains_chinese(self, text: str) -> bool:
        """检测文本是否包含中文"""
        import re
        return bool(re.search(r'[\u4e00-\u9fff]', text))

    def _get_chinese_prompt(self, text: str) -> str:
        """中文简历专用prompt"""
        return f"""
    请仔细分析以下简历内容，准确提取所有信息，以JSON格式返回：

    {text}

    请按以下格式提取（如果某项信息不存在，请设为空值而不是省略字段）：

    {{
    "personal_info": {{
        "full_name": "完整姓名",
        "email": "邮箱地址", 
        "phone": "联系电话",
        "location": "所在地址",
        "linkedin": "LinkedIn链接",
        "github": "GitHub链接", 
        "portfolio": "作品集链接",
        "years_of_experience": "工作年限"
    }},
    "technical_skills": {{
        "programming_languages": ["编程语言列表"],
        "frameworks_libraries": ["框架和库"],
        "databases": ["数据库技术"],
        "cloud_platforms": ["云平台"],
        "tools_software": ["开发工具"],
        "methodologies": ["开发方法论"],
        "soft_skills": ["软技能"]
    }},
    "work_experience": [
        {{
        "job_title": "职位名称",
        "company": "公司名称",
        "location": "工作地点",
        "start_date": "开始日期",
        "end_date": "结束日期", 
        "duration": "工作时长",
        "responsibilities": ["工作职责列表"],
        "achievements": ["工作成就列表"],
        "technologies_used": ["使用的技术"]
        }}
    ],
    "education": [
        {{
        "school": "学校名称",
        "degree": "学位",
        "major": "专业",
        "start_date": "入学时间",
        "end_date": "毕业时间",
        "gpa": "成绩",
        "honors": "荣誉",
        "relevant_courses": ["相关课程"]
        }}
    ],
    "projects": [
        {{
        "name": "项目名称",
        "description": "项目描述",
        "role": "担任角色",
        "duration": "项目时长",
        "team_size": "团队规模",
        "technologies": ["使用技术"],
        "achievements": ["项目成果"],
        "github_url": "代码链接",
        "demo_url": "演示链接"
        }}
    ],
    "certifications": ["证书列表"],
    "notable_achievements": ["重要成就"],
    "languages": ["语言能力"]
    }}

    重要提示：
    1. 仔细查找个人信息，姓名通常在开头，电话是数字格式，邮箱包含@符号
    2. 提取所有工作经历，注意公司名称和职位
    3. 提取教育背景，包括学校和专业
    4. 提取项目经验和技术技能
    5. 如果信息不存在，使用空值而不是省略字段
    6. 保持原文的准确性，不要编造信息

    只返回JSON格式，不要其他说明文字。
    """

    def _get_english_prompt(self, text: str) -> str:
        """英文简历专用prompt"""
        return f"""
    Extract structured information from this resume text and return in JSON format:

    {text}

    Return exact JSON format (use empty values if information doesn't exist, don't omit fields):

    {{
    "personal_info": {{
        "full_name": "Full name from resume",
        "email": "Email address", 
        "phone": "Phone number",
        "location": "Location/Address",
        "linkedin": "LinkedIn URL",
        "github": "GitHub URL", 
        "portfolio": "Portfolio URL",
        "years_of_experience": "Years of experience"
    }},
    "technical_skills": {{
        "programming_languages": ["Programming languages list"],
        "frameworks_libraries": ["Frameworks and libraries"],
        "databases": ["Database technologies"],
        "cloud_platforms": ["Cloud platforms"],
        "tools_software": ["Development tools"],
        "methodologies": ["Development methodologies"],
        "soft_skills": ["Soft skills"]
    }},
    "work_experience": [
        {{
        "job_title": "Job title",
        "company": "Company name",
        "location": "Work location",
        "start_date": "Start date",
        "end_date": "End date", 
        "duration": "Duration",
        "responsibilities": ["List of responsibilities"],
        "achievements": ["List of achievements"],
        "technologies_used": ["Technologies used"]
        }}
    ],
    "education": [
        {{
        "school": "School name",
        "degree": "Degree",
        "major": "Major/Field",
        "start_date": "Start date",
        "end_date": "End date",
        "gpa": "GPA",
        "honors": "Honors",
        "relevant_courses": ["Relevant courses"]
        }}
    ],
    "projects": [
        {{
        "name": "Project name",
        "description": "Project description",
        "role": "Your role",
        "duration": "Project duration",
        "team_size": "Team size",
        "technologies": ["Technologies used"],
        "achievements": ["Project achievements"],
        "github_url": "GitHub URL",
        "demo_url": "Demo URL"
        }}
    ],
    "certifications": ["Certifications list"],
    "notable_achievements": ["Notable achievements"],
    "languages": ["Languages"]
    }}

    Important:
    1. Extract ALL personal information carefully
    2. Extract ALL work experiences with complete details
    3. Extract education background
    4. Extract project experiences
    5. Use empty values if information doesn't exist, don't omit fields
    6. Preserve accuracy from original text

    Return only JSON, no other text.
    """
    
    def _extract_skills_with_regex(self, text: str) -> Dict[str, List[str]]:
        """使用正则表达式辅助提取技术技能"""
        text_lower = text.lower()
        found_skills = {
            'programming_languages': [],
            'frameworks_libraries': [],
            'databases': [],
            'cloud_platforms': [],
            'tools_software': [],
            'methodologies': [],
            'soft_skills': []
        }
        
        # 技术技能提取
        for category, keywords in self.tech_keywords.items():
            for keyword in keywords:
                pattern = r'\b' + re.escape(keyword.lower()) + r'\b'
                if re.search(pattern, text_lower):
                    original_match = self._find_original_case(text, keyword)
                    if original_match and original_match not in found_skills[category]:
                        found_skills[category].append(original_match)
        
        # 软技能提取
        for skill in self.soft_skills_keywords:
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                original_match = self._find_original_case(text, skill)
                if original_match and original_match not in found_skills['soft_skills']:
                    found_skills['soft_skills'].append(original_match)
        
        return found_skills
    
    def _find_original_case(self, text: str, keyword: str) -> str:
        """查找关键词在原文中的大小写格式"""
        pattern = r'\b' + re.escape(keyword.lower()) + r'\b'
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return text[match.start():match.end()]
        return keyword.title()
    
    def _merge_and_validate_results(self, ai_result: Dict, regex_result: Dict) -> Dict:
        """合并AI和正则表达式的结果"""
        final_result = ai_result.copy()
        
        if 'technical_skills' not in final_result:
            final_result['technical_skills'] = {}
        
        for category, skills in regex_result.items():
            if category not in final_result['technical_skills']:
                final_result['technical_skills'][category] = []
            
            existing_skills = [s.lower() for s in final_result['technical_skills'][category]]
            for skill in skills:
                if skill.lower() not in existing_skills:
                    final_result['technical_skills'][category].append(skill)
        
        return final_result
    
    def _calculate_parsing_confidence(self, data: Dict, text: str) -> float:
        """计算解析置信度"""
        confidence_score = 0.0
        max_score = 100.0
        
        # 个人信息完整度 (30分)
        personal_info = data.get('personal_info', {})
        if personal_info.get('full_name'):
            confidence_score += 10
        if personal_info.get('email'):
            confidence_score += 10
        if personal_info.get('phone'):
            confidence_score += 10
        
        # 技术技能提取度 (30分)
        tech_skills = data.get('technical_skills', {})
        total_skills = sum(len(skills) for skills in tech_skills.values() if isinstance(skills, list))
        if total_skills > 0:
            confidence_score += min(30, total_skills * 2)
        
        # 工作经验完整度 (25分)
        work_exp = data.get('work_experience', [])
        if work_exp:
            confidence_score += min(25, len(work_exp) * 8)
        
        # 教育背景完整度 (15分)
        education = data.get('education', [])
        if education:
            confidence_score += min(15, len(education) * 7)
        
        # 文本长度加权
        if len(text.strip()) > 200:
            confidence_score += 10
        
        return min(confidence_score / max_score, 1.0)
    
    def _build_enhanced_resume(self, data: Dict, raw_text: str, confidence: float) -> EnhancedResume:
        """构建增强版简历对象"""
        try:
            personal_data = data.get('personal_info', {})
            personal_info = EnhancedPersonalInfo(**personal_data)
            
            skills_data = data.get('technical_skills', {})
            technical_skills = TechnicalSkills(**skills_data)
            
            work_data = data.get('work_experience', [])
            work_experience = [EnhancedJob(**job) for job in work_data]
            
            edu_data = data.get('education', [])
            education = [EnhancedEducation(**edu) for edu in edu_data]
            
            proj_data = data.get('projects', [])
            projects = [ProjectExperience(**proj) for proj in proj_data]
            
            return EnhancedResume(
                personal_info=personal_info,
                technical_skills=technical_skills,
                work_experience=work_experience,
                education=education,
                projects=projects,
                certifications=data.get('certifications', []),
                notable_achievements=data.get('notable_achievements', []),
                languages=data.get('languages', []),
                raw_text=raw_text,
                parsing_confidence=confidence
            )
            
        except Exception as e:
            logger.error(f"构建简历对象失败: {e}")
            return self._create_fallback_resume(raw_text)
    
    def _create_fallback_structure(self) -> Dict[str, Any]:
        """创建备用数据结构"""
        return {
            "personal_info": {
                "full_name": "",
                "email": "",
                "phone": "",
                "location": "",
                "linkedin": "",
                "github": "",
                "portfolio": "",
                "years_of_experience": ""
            },
            "technical_skills": {
                "programming_languages": [],
                "frameworks_libraries": [],
                "databases": [],
                "cloud_platforms": [],
                "tools_software": [],
                "methodologies": [],
                "soft_skills": []
            },
            "work_experience": [],
            "education": [],
            "projects": [],
            "certifications": [],
            "notable_achievements": [],
            "languages": []
        }
    
    def _create_fallback_resume(self, raw_text: str) -> EnhancedResume:
        """创建备用简历对象"""
        return EnhancedResume(
            personal_info=EnhancedPersonalInfo(),
            technical_skills=TechnicalSkills(),
            work_experience=[],
            education=[],
            projects=[],
            certifications=[],
            notable_achievements=[],
            languages=[],
            raw_text=raw_text,
            parsing_confidence=0.0
        )

if __name__ == "__main__":
    # 示例用法
    async def main():
        parser = EnhancedResumeParser(openai_api_key="your-openai-api-key")
        resume = await parser.parse_resume(file_path="sample_resume.pdf")
        print(json.dumps(asdict(resume), indent=2, ensure_ascii=False))
    
    asyncio.run(main())